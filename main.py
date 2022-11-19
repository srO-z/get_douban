# coding=utf-8
import requests
from lxml import etree
from docx import Document
from docx.shared import Inches
import docx


# from docx.shared import Pt


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/107.0.0.0 Safari/537.36 Edg/107.0.1418.35'}
# 创建docx对象
docu = Document()
# 加入标题
docu.add_heading('豆瓣电影 Top 250', 0)
x = 1
y = 0
for j in range(10):
    url = f'https://movie.douban.com/top250?start={25 * j}&filter='
    resp = requests.get(url, headers=headers)
    e = etree.HTML(resp.text)
    for i in range(25):
        name = e.xpath(f'/html/body/div[3]/div[1]/div/div[1]/ol/li[{i + 1}]/div/div[2]/div[1]/a/span[1]/text()')
        href = e.xpath(f'/html/body/div[3]/div[1]/div/div[1]/ol/li[{i + 1}]/div/div[2]/div[1]/a')
        area = e.xpath(f'/html/body/div[3]/div[1]/div/div[1]/ol/li[{i + 1}]/div/div[2]/div[2]/p[1]/text()[2]')
        quote = e.xpath(f'/html/body/div[3]/div[1]/div/div[1]/ol/li[{i + 1}]/div/div[2]/div[2]/p[2]/span/text()')
        num = e.xpath(f'/html/body/div[3]/div[1]/div/div[1]/ol/li[{i + 1}]/div/div[2]/div[2]/div/span[4]/text()')
        actors = e.xpath(f'/html/body/div[3]/div[1]/div/div[1]/ol/li[{i + 1}]/div/div[2]/div[2]/p[1]/text()[1]')
        # pic_data = e.xpath(f'/html/body/div[3]/div[1]/div/div[1]/ol/li[{i + 1}]/div/div[1]/a/img/@src')

        if y == 2:
            y = 0
            docu.add_page_break()
        name = "《" + str(name).replace('xa0', ',').replace("['", '').replace("']", '') + "》"
        docu.add_heading(f"第{x}部：" + name, 2)

        # data_str = str(pic_data)
        # data_url = data_str.replace("['", '').replace("']", '')
        # resp_pic = requests.get(data_url, headers=headers)
        # with open(f'C:/Users/HuanHuan/Desktop/code/Python/Item/DouBan/IMG/{name}.jpg', 'wb') as f:
        #     f.write(resp_pic.content)
        #     f.close()
        docu.add_picture(f"C:/Users/HuanHuan/Desktop/code/Python/Item/DouBan/IMG/{name}.jpg", width=Inches(1.25))

        actors = "".join(actors)
        actors = actors.split()
        actors = str(actors)
        actors = actors.replace("['", '').replace("']", '').replace("'", '').replace(',', ' ')
        docu.add_paragraph(actors, style='List Bullet')

        area = str(area).replace('xa0', '').replace('\\', '').replace('/', '').replace('n', '').replace(' ',
                                                                                                        '').replace(
            "['",
            '').replace(
            "']", '')
        docu.add_paragraph(area, style='List Bullet')

        quote = str(quote).replace("['", '').replace("']", '')
        docu.add_paragraph(quote, style='List Bullet')

        docu.add_paragraph(num, style='List Bullet')

        href = href[0].attrib
        href = href['href']
        p = docu.add_paragraph()
        if x < 3:
            add_hyperlink(p, href, '按Ctrl+鼠标左键点击这里进入豆瓣网址', 'FF8822', False)
        else:
            add_hyperlink(p, href, '进入豆瓣网址', 'FF8822', False)
        print(f"写入第{x}成功")

        x += 1
        y += 1
docu.save('demo.docx')
print("\n数据下载完毕！")

# # 以下为测试代码
# docu = Document()
# # 加入标题
# docu.add_heading('豆瓣电影 Top 250', 0)
# # 获取网址响应
# url = f'https://movie.douban.com/top250?start=0&filter='
# resp = requests.get(url, headers=headers)
# e = etree.HTML(resp.content)
# name = e.xpath(f'/html/body/div[3]/div[1]/div/div[1]/ol/li[1]/div/div[2]/div[1]/a/span[1]/text()')
# name = "《" + str(name).replace('xa0', ',').replace("['", '').replace("']", '') + "》"
# docu.add_heading(name, 2)
# pic_data = e.xpath('/html/body/div[3]/div[1]/div/div[1]/ol/li[1]/div/div[1]/a/img/@src')
# data_str = str(pic_data)
# data_url = data_str.replace("['", '').replace("']", '')
# resp_pic = requests.get(data_url, headers=headers)
# with open(f'C:/Users/HuanHuan/Desktop/code/Python/Item/DouBan/IMG/{name}.jpg', 'wb') as f:
#     f.write(resp_pic.content)
#     f.close()
# docu.add_picture(f"C:/Users/HuanHuan/Desktop/code/Python/Item/DouBan/IMG/{name}.jpg", width=Inches(1.25))
#
# actors = e.xpath(f'/html/body/div[3]/div[1]/div/div[1]/ol/li[1]/div/div[2]/div[2]/p[1]/text()[1]')
# actors = "".join(actors)
# actors = actors.split()
# actors = str(actors)
# actors = actors.replace("['", '').replace("']", '').replace("'", '').replace(',', ' ')
# docu.add_paragraph(actors, style='List Bullet')
#
# href = e.xpath(f'/html/body/div[3]/div[1]/div/div[1]/ol/li[1]/div/div[2]/div[1]/a')
# href = href[0].attrib
# href = href['href']
# p = docu.add_paragraph()
# add_hyperlink(p, href, '按Ctrl+鼠标左键点击这里进入豆瓣网址', None, True)
#
# area = e.xpath(f'/html/body/div[3]/div[1]/div/div[1]/ol/li[1]/div/div[2]/div[2]/p[1]/text()[2]')
# area = str(area).replace('xa0', '').replace('\\', '').replace('/', '').replace('n', '').replace(' ', '').replace("['",
#                                                                                                                  '').replace(
#     "']", '')
# docu.add_paragraph(area, style='List Bullet')
# quote = e.xpath(f'/html/body/div[3]/div[1]/div/div[1]/ol/li[1]/div/div[2]/div[2]/p[2]/span/text()')
# quote = str(quote).replace("['", '').replace("']", '')
# docu.add_paragraph(quote, style='List Bullet')
#
# num = e.xpath(f'/html/body/div[3]/div[1]/div/div[1]/ol/li[1]/div/div[2]/div[2]/div/span[4]/text()')
# docu.add_paragraph(num, style='List Bullet')
#
# docu.save('demo.docx')
# print("保存成功！")
